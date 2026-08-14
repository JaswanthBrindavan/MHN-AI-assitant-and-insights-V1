"""Ingest scripts — drug CSV (scripts.ingest_drugs) and MCP docx corpus
(scripts.ingest_mcp_corpus).

CSV fixtures are written to tmp_path with the real merged-database header
subset and a UTF-8 BOM on the first line, mirroring the production file.
Docx fixtures are built with python-docx (title, AKA line, Definition
paragraph, Symptoms table) so the full parse → chunk → upsert path runs.
"""

from __future__ import annotations

import csv
import io
from pathlib import Path

import pytest
from docx import Document as DocxDocument
from sqlalchemy import func, select

import scripts.ingest_drugs as ingest_drugs
from app.models.chat import McpChunk
from app.models.knowledge import ConditionRegistry, DrugReference
from scripts.ingest_drugs import ingest_drug_csv, row_to_record
from scripts.ingest_mcp_corpus import ingest_mcp_folder


def rec_of(row: dict) -> dict:
    """row_to_record with a non-None narrowing assert (typing helper)."""
    rec = row_to_record(row)
    assert rec is not None
    return rec

UTF8_BOM = b"\xef\xbb\xbf"

DRUG_HEADER = [
    "name", "price(₹)", "Is_discontinued", "manufacturer_name", "type",
    "pack_size_label", "short_composition1", "short_composition2",
    "sideEffect0", "sideEffect1", "sideEffect2", "sideEffect3",
    "use0", "use1", "use2", "substitute0", "substitute1",
    "Chemical Class", "Habit Forming", "Therapeutic Class", "Action Class",
    "id_indian",
]


def _write_drug_csv(path: Path, rows: list[dict]) -> None:
    """Write a CSV with the real header subset and a leading UTF-8 BOM."""
    buf = io.StringIO()
    writer = csv.DictWriter(buf, fieldnames=DRUG_HEADER, restval="", extrasaction="ignore")
    writer.writeheader()
    writer.writerows(rows)
    path.write_bytes(UTF8_BOM + buf.getvalue().encode("utf-8"))


async def _count(db, model) -> int:
    return (await db.execute(select(func.count()).select_from(model))).scalar_one()


# --------------------------------------------------------------------------- #
# row_to_record — name handling
# --------------------------------------------------------------------------- #
def test_blank_name_returns_none():
    assert row_to_record({"name": ""}) is None


def test_whitespace_only_name_returns_none():
    assert row_to_record({"name": "   \t  "}) is None


def test_missing_name_key_returns_none():
    assert row_to_record({}) is None


def test_minimal_row_defaults():
    rec = rec_of({"name": "Dolo 650 Tablet"})
    assert rec is not None
    assert rec["name"] == "Dolo 650 Tablet"
    assert rec["name_normalized"] == "dolo 650 tablet"
    assert rec["source_id"] is None
    assert rec["manufacturer"] is None
    assert rec["dosage_type"] is None
    assert rec["pack_size"] is None
    assert rec["price_inr"] is None
    assert rec["is_discontinued"] is False
    assert rec["composition1"] is None
    assert rec["composition2"] is None
    assert rec["composition_normalized"] is None
    assert rec["side_effects"] is None
    assert rec["uses"] is None
    assert rec["substitutes"] is None
    assert rec["chemical_class"] is None
    assert rec["habit_forming"] is None
    assert rec["therapeutic_class"] is None
    assert rec["action_class"] is None


def test_name_normalized_collapses_whitespace_and_lowercases():
    rec = rec_of({"name": "  Augmentin   625  DUO\tTablet "})
    assert rec["name"] == "Augmentin   625  DUO\tTablet"
    assert rec["name_normalized"] == "augmentin 625 duo tablet"


def test_unicode_name_preserved():
    rec = rec_of({"name": "Paracétamol® 500 mg गोली"})
    assert rec["name"] == "Paracétamol® 500 mg गोली"
    assert rec["name_normalized"] == "paracétamol® 500 mg गोली"


def test_long_name_clipped_to_255():
    rec = rec_of({"name": "X" * 300})
    assert len(rec["name"]) == 255
    assert len(rec["name_normalized"]) == 255


# --------------------------------------------------------------------------- #
# row_to_record — price parsing
# --------------------------------------------------------------------------- #
@pytest.mark.parametrize(
    ("raw", "expected"),
    [
        ("NA", None),
        ("na", None),
        ("", None),
        ("   ", None),
        ("garbage", None),
        ("₹45", None),
        ("45,50", None),
        ("45.5", 45.5),
        (" 45.5 ", 45.5),
        ("0", 0.0),
        ("1e3", 1000.0),
    ],
)
def test_price_parsing(raw, expected):
    rec = rec_of({"name": "N", "price(₹)": raw})
    assert rec["price_inr"] == expected


def test_price_fallback_plain_price_column():
    rec = rec_of({"name": "N", "price": "12.5"})
    assert rec["price_inr"] == 12.5


# --------------------------------------------------------------------------- #
# row_to_record — Is_discontinued parsing
# --------------------------------------------------------------------------- #
@pytest.mark.parametrize(
    ("raw", "expected"),
    [
        ("TRUE", True),
        ("true", True),
        (" True ", True),
        ("1", True),
        ("yes", True),
        ("False", False),
        ("FALSE", False),
        ("", False),
        ("0", False),
        ("no", False),
        ("discontinued", False),
    ],
)
def test_is_discontinued_parsing(raw, expected):
    rec = rec_of({"name": "N", "Is_discontinued": raw})
    assert rec["is_discontinued"] is expected


# --------------------------------------------------------------------------- #
# row_to_record — side effects / uses / substitutes collection
# --------------------------------------------------------------------------- #
def test_side_effects_deduped_and_na_filtered():
    rec = rec_of(
        {
            "name": "N",
            "sideEffect0": "Nausea",
            "sideEffect1": "Nausea",
            "sideEffect2": "NA",
            "sideEffect3": "na",
            "sideEffect4": "Headache",
        }
    )
    assert rec["side_effects"] == ["Nausea", "Headache"]


def test_side_effects_gap_in_indices_still_collected():
    rec = rec_of({"name": "N", "sideEffect0": "", "sideEffect3": "Dizziness"})
    assert rec["side_effects"] == ["Dizziness"]


def test_side_effects_case_variants_not_deduped():
    # Dedup is exact-match: case variants are both kept (current behavior).
    rec = rec_of({"name": "N", "sideEffect0": "Nausea", "sideEffect1": "nausea"})
    assert rec["side_effects"] == ["Nausea", "nausea"]


def test_side_effects_index_41_is_last_collected():
    rec = rec_of({"name": "N", "sideEffect41": "Rash", "sideEffect42": "Ignored"})
    assert rec["side_effects"] == ["Rash"]


def test_all_na_side_effects_becomes_none():
    rec = rec_of({"name": "N", "sideEffect0": "NA", "sideEffect1": "NA"})
    assert rec["side_effects"] is None


def test_uses_collected_in_order_max_five():
    row = {"name": "N"}
    row.update({f"use{i}": f"Use {i}" for i in range(6)})
    rec = rec_of(row)
    assert rec["uses"] == [f"Use {i}" for i in range(5)]  # use5 outside range


def test_substitutes_collected_and_deduped():
    rec = rec_of(
        {
            "name": "N",
            "substitute0": "Calpol 500",
            "substitute1": "NA",
            "substitute2": "Calpol 500",
            "substitute3": "Pacimol 500",
        }
    )
    assert rec["substitutes"] == ["Calpol 500", "Pacimol 500"]


# --------------------------------------------------------------------------- #
# row_to_record — compositions, classes, source id, clipping
# --------------------------------------------------------------------------- #
def test_composition_normalized_joins_and_lowercases():
    rec = rec_of(
        {
            "name": "N",
            "short_composition1": " Amoxycillin  (500mg) ",
            "short_composition2": "Clavulanic Acid (125mg)",
        }
    )
    assert rec["composition1"] == "Amoxycillin  (500mg)"
    assert rec["composition2"] == "Clavulanic Acid (125mg)"
    assert rec["composition_normalized"] == "amoxycillin (500mg) clavulanic acid (125mg)"


def test_composition_normalized_single_component():
    rec = rec_of({"name": "N", "short_composition1": "Metformin (500mg)"})
    assert rec["composition2"] is None
    assert rec["composition_normalized"] == "metformin (500mg)"


def test_composition_normalized_clipped_to_512():
    rec = rec_of(
        {"name": "N", "short_composition1": "A" * 300, "short_composition2": "B" * 300}
    )
    assert len(rec["composition1"]) == 255
    assert len(rec["composition_normalized"]) == 512
    assert rec["composition_normalized"] == ("a" * 300 + " " + "b" * 300)[:512]


def test_habit_forming_na_is_none():
    for raw in ("NA", "na", "Na", ""):
        assert rec_of({"name": "N", "Habit Forming": raw})["habit_forming"] is None


def test_habit_forming_value_kept_and_clipped_to_16():
    rec = rec_of({"name": "N", "Habit Forming": "No"})
    assert rec["habit_forming"] == "No"
    rec = rec_of({"name": "N", "Habit Forming": "Y" * 20})
    assert rec["habit_forming"] == "Y" * 16


def test_source_id_prefers_id_indian_and_clips_to_32():
    rec = rec_of({"name": "N", "id_indian": "abc123", "id_dataset": "zzz"})
    assert rec["source_id"] == "abc123"
    rec = rec_of({"name": "N", "id_indian": "", "id_dataset": "D" * 40})
    assert rec["source_id"] == "D" * 32


def test_classes_trimmed_and_clipped():
    rec = rec_of(
        {
            "name": "N",
            "Chemical Class": " Biguanides ",
            "Therapeutic Class": "T" * 200,
            "Action Class": "A" * 200,
            "type": "allopathy" * 20,
            "pack_size_label": "P" * 200,
            "manufacturer_name": "M" * 300,
        }
    )
    assert rec["chemical_class"] == "Biguanides"
    assert len(rec["therapeutic_class"]) == 128
    assert len(rec["action_class"]) == 128
    assert len(rec["dosage_type"]) == 64
    assert len(rec["pack_size"]) == 128
    assert len(rec["manufacturer"]) == 255


# --------------------------------------------------------------------------- #
# ingest_drug_csv — end-to-end against the DB
# --------------------------------------------------------------------------- #
async def test_ingest_csv_bom_handled_and_rows_inserted(db_session, tmp_path):
    path = tmp_path / "drugs.csv"
    _write_drug_csv(
        path,
        [
            {"name": "Dolo 650 Tablet", "price(₹)": "30.91", "manufacturer_name": "Micro Labs"},
            {"name": "Crocin Advance", "price(₹)": "NA", "Is_discontinued": "TRUE"},
        ],
    )
    assert path.read_bytes().startswith(UTF8_BOM)  # fixture really has a BOM

    stats = await ingest_drug_csv(db_session, path)
    assert stats == {"rows": 2, "inserted": 2, "skipped": 0}

    rows = (await db_session.execute(select(DrugReference).order_by(DrugReference.name)))
    rows = rows.scalars().all()
    # If the BOM leaked into the header the "name" key would not match and
    # every row would be skipped — reaching here proves it was stripped.
    assert [r.name for r in rows] == ["Crocin Advance", "Dolo 650 Tablet"]
    dolo = rows[1]
    assert dolo.price_inr == 30.91
    assert dolo.manufacturer == "Micro Labs"
    assert dolo.is_discontinued is False
    assert rows[0].price_inr is None
    assert rows[0].is_discontinued is True


async def test_ingest_csv_blank_name_rows_skipped_and_counted(db_session, tmp_path):
    path = tmp_path / "drugs.csv"
    _write_drug_csv(
        path,
        [
            {"name": "Valid Drug", "price(₹)": "10"},
            {"name": "", "price(₹)": "20"},
            {"name": "   ", "price(₹)": "30"},
        ],
    )
    stats = await ingest_drug_csv(db_session, path)
    assert stats == {"rows": 3, "inserted": 1, "skipped": 2}
    assert await _count(db_session, DrugReference) == 1


async def test_ingest_csv_quoted_field_with_commas(db_session, tmp_path):
    path = tmp_path / "drugs.csv"
    name = "Combiflam, Extra (400mg, 325mg)"
    _write_drug_csv(
        path,
        [
            {
                "name": name,
                "short_composition1": "Ibuprofen (400mg), Paracetamol (325mg)",
                "sideEffect0": "Nausea, mild",
            }
        ],
    )
    # The comma-bearing fields must have been quoted by the writer.
    assert b'"' in path.read_bytes()

    stats = await ingest_drug_csv(db_session, path)
    assert stats["inserted"] == 1
    row = (await db_session.execute(select(DrugReference))).scalars().one()
    assert row.name == name
    assert row.composition1 == "Ibuprofen (400mg), Paracetamol (325mg)"
    assert row.side_effects == ["Nausea, mild"]


async def test_ingest_csv_batches_across_batch_size_boundary(db_session, tmp_path, monkeypatch):
    monkeypatch.setattr(ingest_drugs, "BATCH_SIZE", 3)
    path = tmp_path / "drugs.csv"
    _write_drug_csv(path, [{"name": f"Drug {i}", "price(₹)": str(i)} for i in range(7)])

    stats = await ingest_drug_csv(db_session, path)
    assert stats == {"rows": 7, "inserted": 7, "skipped": 0}
    assert await _count(db_session, DrugReference) == 7
    names = (await db_session.execute(select(DrugReference.name))).scalars().all()
    assert sorted(names) == [f"Drug {i}" for i in range(7)]


async def test_ingest_csv_exact_batch_multiple(db_session, tmp_path, monkeypatch):
    monkeypatch.setattr(ingest_drugs, "BATCH_SIZE", 3)
    path = tmp_path / "drugs.csv"
    _write_drug_csv(path, [{"name": f"Drug {i}"} for i in range(6)])
    stats = await ingest_drug_csv(db_session, path)
    assert stats["inserted"] == 6
    assert await _count(db_session, DrugReference) == 6


async def test_ingest_csv_reingest_same_file_row_count_stable(db_session, tmp_path):
    path = tmp_path / "drugs.csv"
    _write_drug_csv(path, [{"name": f"Drug {i}"} for i in range(4)])

    first = await ingest_drug_csv(db_session, path)
    assert first["inserted"] == 4
    assert await _count(db_session, DrugReference) == 4

    second = await ingest_drug_csv(db_session, path)
    assert second["inserted"] == 4
    assert await _count(db_session, DrugReference) == 4  # truncated, not doubled


async def test_ingest_csv_reingest_replaces_previous_contents(db_session, tmp_path):
    old = tmp_path / "old.csv"
    new = tmp_path / "new.csv"
    _write_drug_csv(old, [{"name": "Old A"}, {"name": "Old B"}, {"name": "Old C"}])
    _write_drug_csv(new, [{"name": "New A"}, {"name": "New B"}])

    await ingest_drug_csv(db_session, old)
    await ingest_drug_csv(db_session, new)
    names = (await db_session.execute(select(DrugReference.name))).scalars().all()
    assert sorted(names) == ["New A", "New B"]


async def test_ingest_csv_empty_file_no_rows(db_session, tmp_path):
    path = tmp_path / "drugs.csv"
    _write_drug_csv(path, [])
    stats = await ingest_drug_csv(db_session, path)
    assert stats == {"rows": 0, "inserted": 0, "skipped": 0}
    assert await _count(db_session, DrugReference) == 0


async def test_ingest_csv_unicode_values_roundtrip(db_session, tmp_path):
    path = tmp_path / "drugs.csv"
    _write_drug_csv(
        path,
        [
            {
                "name": "Paracétamol गोली 錠剤",
                "price(₹)": "12.5",
                "manufacturer_name": "Ærøskøbing Pharma",
                "sideEffect0": "जी मिचलाना",
            }
        ],
    )
    stats = await ingest_drug_csv(db_session, path)
    assert stats["inserted"] == 1
    row = (await db_session.execute(select(DrugReference))).scalars().one()
    assert row.name == "Paracétamol गोली 錠剤"
    assert row.name_normalized == "paracétamol गोली 錠剤"
    assert row.manufacturer == "Ærøskøbing Pharma"
    assert row.side_effects == ["जी मिचलाना"]
    assert row.price_inr == 12.5


# --------------------------------------------------------------------------- #
# MCP corpus fixtures
# --------------------------------------------------------------------------- #
def _write_mcp_docx(
    path: Path,
    title: str,
    aka: str | None = "AKA i.e., Sugar Disease, NIDD",
    definition: str | None = "A chronic metabolic disorder of glucose regulation.",
    with_symptoms_table: bool = True,
) -> None:
    """Tiny but structurally faithful MCP docx: title, AKA, two sections."""
    doc = DocxDocument()
    doc.add_paragraph(title)
    if aka:
        doc.add_paragraph(aka)
    if definition is not None:
        doc.add_paragraph("Definition")
        doc.add_paragraph(definition)
    if with_symptoms_table:
        doc.add_paragraph("Symptoms")
        table = doc.add_table(rows=3, cols=2)
        cells = [
            ("Symptom", "Notes"),
            ("Thirst", "Increased"),
            ("Fatigue", "Common"),
        ]
        for r, (a, b) in enumerate(cells):
            table.rows[r].cells[0].text = a
            table.rows[r].cells[1].text = b
    doc.save(str(path))


async def _chunk_count(db, code: str) -> int:
    stmt = select(func.count()).select_from(McpChunk).where(McpChunk.condition_code == code)
    return (await db.execute(stmt)).scalar_one()


# --------------------------------------------------------------------------- #
# ingest_mcp_folder
# --------------------------------------------------------------------------- #
async def test_mcp_single_file_full_ingest(db_session, tmp_path):
    _write_mcp_docx(tmp_path / "MC001_Diabetes_Mellitus.docx", "Diabetes Mellitus Type 2")
    stats = await ingest_mcp_folder(db_session, tmp_path)

    assert stats["files"] == 1
    assert stats["ingested"] == 1
    assert stats["chunks"] == 2  # definition + symptoms
    assert stats["duplicates"] == []
    assert stats["errors"] == []

    row = (await db_session.execute(select(ConditionRegistry))).scalars().one()
    assert row.condition_code == "MC001"
    assert row.display_name == "Diabetes Mellitus Type 2"
    assert row.aliases == ["Sugar Disease", "NIDD"]
    assert row.engine_codes == ["T2DM"]  # ENGINE_CODE_MAP applied for MC001
    assert row.source_file == "MC001_Diabetes_Mellitus.docx"
    assert row.active is True

    chunks = (await db_session.execute(select(McpChunk))).scalars().all()
    assert {c.chunk_type for c in chunks} == {"definition", "symptoms"}
    by_type = {c.chunk_type: c for c in chunks}
    assert "Diabetes Mellitus Type 2 — definition:" in by_type["definition"].content
    assert "chronic metabolic disorder" in by_type["definition"].content
    assert "Symptom: Thirst; Notes: Increased" in by_type["symptoms"].content
    assert all(c.condition_code == "MC001" for c in chunks)
    assert all(c.embedding is None for c in chunks)
    meta = by_type["definition"].chunk_metadata
    assert meta["source"] == "mcp_master_profile"
    assert meta["source_file"] == "MC001_Diabetes_Mellitus.docx"
    assert meta["display_name"] == "Diabetes Mellitus Type 2"


async def test_mcp_unmapped_code_gets_empty_engine_codes(db_session, tmp_path):
    _write_mcp_docx(tmp_path / "MC305_Alcohol_Use.docx", "Alcohol Use Disorder")
    await ingest_mcp_folder(db_session, tmp_path)
    row = (await db_session.execute(select(ConditionRegistry))).scalars().one()
    assert row.condition_code == "MC305"
    assert row.engine_codes == []


async def test_mcp_duplicate_code_second_file_reported(db_session, tmp_path):
    _write_mcp_docx(tmp_path / "MC305_First.docx", "Alcohol Use Disorder")
    _write_mcp_docx(tmp_path / "MC305_Second.docx", "Duplicate Profile")
    stats = await ingest_mcp_folder(db_session, tmp_path)

    assert stats["files"] == 2
    assert stats["ingested"] == 1
    assert stats["duplicates"] == ["MC305_Second.docx"]  # first in sorted order wins
    assert stats["errors"] == []

    row = (await db_session.execute(select(ConditionRegistry))).scalars().one()
    assert row.display_name == "Alcohol Use Disorder"
    assert row.source_file == "MC305_First.docx"
    assert await _chunk_count(db_session, "MC305") == 2


async def test_mcp_corrupt_file_reported_others_ingested(db_session, tmp_path):
    _write_mcp_docx(tmp_path / "MC001_Diabetes.docx", "Diabetes Mellitus Type 2")
    (tmp_path / "MC999_Corrupt.docx").write_text("this is not a zip archive")
    stats = await ingest_mcp_folder(db_session, tmp_path)

    assert stats["files"] == 2
    assert stats["ingested"] == 1
    assert len(stats["errors"]) == 1
    assert stats["errors"][0].startswith("MC999_Corrupt.docx:")
    assert await _count(db_session, ConditionRegistry) == 1
    assert await _chunk_count(db_session, "MC001") == 2


async def test_mcp_word_lock_file_ignored(db_session, tmp_path):
    _write_mcp_docx(tmp_path / "MC001_Diabetes.docx", "Diabetes Mellitus Type 2")
    # Word lock files are junk bytes; must be skipped before parsing.
    (tmp_path / "~$MC001_Diabetes.docx").write_bytes(b"\x00\x01lockfile")
    stats = await ingest_mcp_folder(db_session, tmp_path)

    assert stats["ingested"] == 1
    assert stats["errors"] == []
    assert stats["duplicates"] == []
    # NOTE(potential-bug): stats["files"] counts ~$ lock files (glob matches
    # them before the skip) — cosmetic inaccuracy in the printed summary.
    assert stats["files"] == 2


async def test_mcp_reingest_replaces_chunks_no_duplicates(db_session, tmp_path):
    _write_mcp_docx(tmp_path / "MC001_Diabetes.docx", "Diabetes Mellitus Type 2")
    first = await ingest_mcp_folder(db_session, tmp_path)
    assert first["chunks"] == 2

    second = await ingest_mcp_folder(db_session, tmp_path)
    assert second["ingested"] == 1
    assert second["chunks"] == 2
    assert await _chunk_count(db_session, "MC001") == 2  # replaced, not appended
    assert await _count(db_session, ConditionRegistry) == 1  # upsert, not insert


async def test_mcp_reingest_updates_registry_row_in_place(db_session, tmp_path):
    path = tmp_path / "MC001_Diabetes.docx"
    _write_mcp_docx(path, "Diabetes Mellitus Type 2")
    await ingest_mcp_folder(db_session, tmp_path)
    original_id = (await db_session.execute(select(ConditionRegistry))).scalars().one().id

    _write_mcp_docx(path, "Diabetes Mellitus (Type II)", aka="AKA i.e., T2D")
    await ingest_mcp_folder(db_session, tmp_path)

    row = (await db_session.execute(select(ConditionRegistry))).scalars().one()
    assert row.id == original_id  # same row, updated in place
    assert row.display_name == "Diabetes Mellitus (Type II)"
    assert row.aliases == ["T2D"]
    chunk = (
        await db_session.execute(select(McpChunk).where(McpChunk.chunk_type == "definition"))
    ).scalars().one()
    assert chunk.content.startswith("Diabetes Mellitus (Type II) — definition:")


async def test_mcp_file_without_code_in_name_is_error(db_session, tmp_path):
    _write_mcp_docx(tmp_path / "Untitled_Profile.docx", "Some Condition")
    stats = await ingest_mcp_folder(db_session, tmp_path)
    assert stats["ingested"] == 0
    assert stats["errors"] == ["Untitled_Profile.docx: missing MC code or title"]
    assert await _count(db_session, ConditionRegistry) == 0


async def test_mcp_title_only_file_produces_no_chunks_error(db_session, tmp_path):
    _write_mcp_docx(
        tmp_path / "MC777_Empty.docx",
        "Empty Profile",
        aka=None,
        definition=None,
        with_symptoms_table=False,
    )
    stats = await ingest_mcp_folder(db_session, tmp_path)
    assert stats["ingested"] == 0
    assert stats["errors"] == ["MC777_Empty.docx: produced no chunks"]
    assert await _count(db_session, ConditionRegistry) == 0


async def test_mcp_empty_folder(db_session, tmp_path):
    stats = await ingest_mcp_folder(db_session, tmp_path)
    assert stats == {
        "files": 0, "ingested": 0, "chunks": 0, "duplicates": [],
        "duplicate_content": [], "errors": [],
    }


async def test_mcp_mixed_folder_stats_accurate(db_session, tmp_path):
    _write_mcp_docx(tmp_path / "MC001_Diabetes.docx", "Diabetes Mellitus Type 2")
    _write_mcp_docx(
        tmp_path / "MC051_Hypertension.docx", "Primary Hypertension", aka="AKA i.e., High BP"
    )
    _write_mcp_docx(tmp_path / "MC305_First.docx", "Alcohol Use Disorder")
    _write_mcp_docx(tmp_path / "MC305_Second.docx", "Duplicate Profile")
    (tmp_path / "MC999_Corrupt.docx").write_text("not a docx")
    (tmp_path / "~$MC001_Diabetes.docx").write_bytes(b"lock")
    _write_mcp_docx(tmp_path / "NoCode.docx", "Anonymous Condition")

    stats = await ingest_mcp_folder(db_session, tmp_path)
    assert stats["files"] == 7  # every *.docx, lock file included
    assert stats["ingested"] == 3  # MC001, MC051, MC305 (first)
    assert stats["chunks"] == 6  # 2 per ingested condition
    assert stats["duplicates"] == ["MC305_Second.docx"]
    error_files = sorted(e.split(":", 1)[0] for e in stats["errors"])
    assert error_files == ["MC999_Corrupt.docx", "NoCode.docx"]

    assert await _count(db_session, ConditionRegistry) == 3
    assert await _count(db_session, McpChunk) == 6

    codes = (
        await db_session.execute(
            select(ConditionRegistry.condition_code).order_by(ConditionRegistry.condition_code)
        )
    ).scalars().all()
    assert codes == ["MC001", "MC051", "MC305"]

    # ENGINE_CODE_MAP applied where defined, empty elsewhere.
    rows = (await db_session.execute(select(ConditionRegistry))).scalars().all()
    engine = {r.condition_code: r.engine_codes for r in rows}
    assert engine == {"MC001": ["T2DM"], "MC051": ["HTN"], "MC305": []}
