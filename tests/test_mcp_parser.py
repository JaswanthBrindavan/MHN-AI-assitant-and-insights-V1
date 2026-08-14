"""Tests for app/knowledge/mcp_parser.py — the MCP docx corpus parser.

Fixture .docx files are built at runtime with python-docx into tmp_path,
modeled on the real corpus shape: title paragraph, "AKA i.e., ..." line,
section-marker paragraphs, and data tables in document order.
"""

from __future__ import annotations

import pytest
from docx import Document

from app.knowledge.mcp_parser import (
    MAX_CHUNK_CHARS,
    SECTION_MARKERS,
    ParsedMcp,
    _split_to_chunks,
    _table_to_lines,
    build_chunks,
    extract_code_from_filename,
    parse_aka_line,
    parse_mcp_docx,
)

# --------------------------------------------------------------------------- #
# Helpers: build corpus-shaped docx fixtures at runtime
# --------------------------------------------------------------------------- #


def _fill(table, rows):
    for r, values in enumerate(rows):
        for c, value in enumerate(values):
            table.cell(r, c).text = value


def _write_corpus_docx(tmp_path, filename="MC001 Hypertension.docx"):
    """A representative MCP document exercising markers, tables, and blanks."""
    doc = Document()
    doc.add_paragraph("Hypertension")
    doc.add_paragraph("AKA i.e., High Blood Pressure, HTN (Htn), Raised BP")
    doc.add_paragraph("Introduction to MCP")
    doc.add_paragraph("This intro text belongs to a structural header.")
    doc.add_paragraph("Definition")
    doc.add_paragraph("A sustained elevation of arterial blood pressure.")
    doc.add_paragraph("")  # blank paragraph: skipped
    doc.add_paragraph("Prevalence")
    doc.add_paragraph("Common in adults worldwide.")
    doc.add_paragraph("Diagnosis")
    doc.add_paragraph("Repeated readings above the threshold on separate days.")
    doc.add_paragraph("Research Draft of MCP")
    doc.add_paragraph("Profiles Prone To")
    t = doc.add_table(rows=3, cols=2)
    _fill(
        t,
        [
            ["Profile", "Reason"],
            ["Sedentary adults", "Low activity"],
            ["Smokers", "Vascular strain"],
        ],
    )
    doc.add_paragraph("Symptoms")
    doc.add_paragraph("Often silent in early stages.")
    t2 = doc.add_table(rows=3, cols=3)
    _fill(
        t2,
        [
            ["Symptom", "Type", "Note"],
            ["Headache", "Common", "Morning pattern"],
            ["Dizziness", "Occasional", ""],
        ],
    )
    doc.add_paragraph("Suggestions")
    doc.add_paragraph("Discuss readings with a clinician.")
    path = tmp_path / filename
    doc.save(str(path))
    return path


def _write_docx(tmp_path, filename, paragraphs, tables_after=None):
    """Write paragraphs in order; tables_after maps paragraph index -> rows."""
    doc = Document()
    tables_after = tables_after or {}
    for i, text in enumerate(paragraphs):
        doc.add_paragraph(text)
        if i in tables_after:
            rows = tables_after[i]
            t = doc.add_table(rows=len(rows), cols=len(rows[0]))
            _fill(t, rows)
    path = tmp_path / filename
    doc.save(str(path))
    return path


def _table(rows_data, cols=None):
    """Build an in-memory docx Table with the given cell text."""
    doc = Document()
    n_cols = cols if cols is not None else (len(rows_data[0]) if rows_data else 1)
    t = doc.add_table(rows=len(rows_data), cols=n_cols)
    _fill(t, rows_data)
    return t


# --------------------------------------------------------------------------- #
# extract_code_from_filename
# --------------------------------------------------------------------------- #


@pytest.mark.parametrize(
    ("filename", "expected"),
    [
        ("MC001 Hypertension.docx", "MC001"),  # spaced
        ("MC305_Alcohol_Use_Disorder.docx", "MC305"),  # underscored
        ("MC1234 Rare Condition.docx", "MC1234"),  # 4-digit
        ("Hypertension Profile.docx", None),  # no code
        ("Notes MC042 final.docx", "MC042"),  # code in middle
        ("MC001.docx", "MC001"),  # code then extension dot
        ("MC9999", "MC9999"),  # code at end of string
        ("MC12 short.docx", None),  # too few digits
        ("MC12345 five digits.docx", None),  # too many digits
        ("XMC001 glued.docx", None),  # no leading word boundary
        ("mc001 lowercase.docx", None),  # case-sensitive
        ("", None),  # empty filename
        ("MCP overview.docx", None),  # MC without digits
    ],
)
def test_extract_code_from_filename(filename, expected):
    assert extract_code_from_filename(filename) == expected


# --------------------------------------------------------------------------- #
# parse_aka_line
# --------------------------------------------------------------------------- #


def test_aka_ie_comma_form():
    line = "AKA i.e., High Blood Pressure, HTN, Raised BP"
    assert parse_aka_line(line) == ["High Blood Pressure", "HTN", "Raised BP"]


def test_aka_colon_form():
    assert parse_aka_line("AKA: Flu, Grippe") == ["Flu", "Grippe"]


def test_aka_lowercase_no_ie():
    assert parse_aka_line("aka influenza") == ["influenza"]


def test_aka_ie_without_trailing_comma():
    assert parse_aka_line("AKA i.e. Influenza, Grippe") == ["Influenza", "Grippe"]


def test_aka_empty_alias_segments_dropped():
    assert parse_aka_line("AKA i.e., A,, B,") == ["A", "B"]


def test_aka_case_insensitive_dedupe_keeps_first_spelling():
    assert parse_aka_line("AKA i.e., Flu, flu, FLU") == ["Flu"]


def test_aka_parenthesised_abbreviation_expanded():
    # "Name (ABBR)" emits the whole part, the paren-free base, and the inner.
    line = "AKA i.e., Polycystic ovary syndrome (PCOS), Stein-Leventhal syndrome"
    assert parse_aka_line(line) == [
        "Polycystic ovary syndrome (PCOS)",
        "Polycystic ovary syndrome",
        "PCOS",
        "Stein-Leventhal syndrome",
    ]


def test_aka_paren_inner_duplicate_of_base_not_repeated():
    # Inner "Htn" collapses (case-insensitively) into the base "HTN".
    assert parse_aka_line("AKA i.e., HTN (Htn), Raised BP") == ["HTN (Htn)", "HTN", "Raised BP"]


def test_aka_commentary_parens_not_emitted_as_alias():
    # "lay/term/clinical/ayurvedic/form" and +/ inners are commentary: only
    # the whole part and the paren-free base are kept.
    assert parse_aka_line("AKA i.e., Flu (lay term), Grippe") == ["Flu (lay term)", "Flu", "Grippe"]
    assert parse_aka_line("AKA i.e., Mix (A/B), Plus (C+D)") == [
        "Mix (A/B)",
        "Mix",
        "Plus (C+D)",
        "Plus",
    ]
    line = "AKA i.e., Hypertension (Ayurvedic form), Raised BP"
    assert parse_aka_line(line) == ["Hypertension (Ayurvedic form)", "Hypertension", "Raised BP"]


def test_aka_paren_inner_length_bounds():
    # Inners shorter than 2 or longer than 40 chars are not extracted.
    assert parse_aka_line("AKA i.e., X (a), Y") == ["X (a)", "X", "Y"]
    long_inner = "z" * 41
    assert parse_aka_line(f"AKA i.e., Long ({long_inner}), Y") == [
        f"Long ({long_inner})",
        "Long",
        "Y",
    ]


def test_aka_semicolon_separators():
    assert parse_aka_line("AKA i.e., A; B, C") == ["A", "B", "C"]


def test_aka_curly_quotes_stripped():
    assert parse_aka_line("AKA: “Sugar disease”, Diabetes") == [
        "Sugar disease",
        "Diabetes",
    ]


def test_aka_no_aliases_after_prefix():
    assert parse_aka_line("AKA i.e.,") == []
    assert parse_aka_line("AKA") == []
    assert parse_aka_line("") == []


def test_aka_whitespace_padding_stripped():
    assert parse_aka_line("   AKA i.e.,   Flu   ") == ["Flu"]


def test_aka_unicode_aliases():
    line = "AKA i.e., Ménière's disease, Labyrinthine hydrops"
    assert parse_aka_line(line) == ["Ménière's disease", "Labyrinthine hydrops"]


def test_aka_trailing_periods_stripped_from_alias():
    # .strip(".") removes the trailing dot of dotted abbreviations too.
    assert parse_aka_line("AKA i.e. B.P., HTN") == ["B.P", "HTN"]


def test_aka_alias_starting_with_ie_loses_prefix():
    # NOTE(potential-bug): the optional "i.e." prefix regex consumes the
    # leading "IE" of a real alias ("IED trauma" -> "D trauma").
    assert parse_aka_line("AKA IED trauma") == ["D trauma"]


# --------------------------------------------------------------------------- #
# parse_mcp_docx: preamble + fields
# --------------------------------------------------------------------------- #


def test_corpus_doc_fields(tmp_path):
    path = _write_corpus_docx(tmp_path)
    parsed = parse_mcp_docx(path)
    assert parsed.code == "MC001"
    assert parsed.display_name == "Hypertension"
    assert parsed.source_file == "MC001 Hypertension.docx"
    assert parsed.aliases == ["High Blood Pressure", "HTN (Htn)", "HTN", "Raised BP"]


def test_no_code_in_filename_yields_unknown(tmp_path):
    path = _write_corpus_docx(tmp_path, filename="Hypertension Master Profile.docx")
    parsed = parse_mcp_docx(path)
    assert parsed.code == "UNKNOWN"
    assert parsed.source_file == "Hypertension Master Profile.docx"


def test_accepts_str_path(tmp_path):
    path = _write_corpus_docx(tmp_path)
    parsed = parse_mcp_docx(str(path))
    assert parsed.display_name == "Hypertension"


def test_doc_without_aka_line(tmp_path):
    path = _write_docx(
        tmp_path,
        "MC010 Plain.docx",
        ["Plain Condition", "Definition", "Some definition text."],
    )
    parsed = parse_mcp_docx(path)
    assert parsed.aliases == []
    assert parsed.sections["definition"] == ["Some definition text."]


def test_unicode_display_name(tmp_path):
    path = _write_docx(
        tmp_path,
        "MC777 Meniere.docx",
        ["Ménière's Disease", "AKA i.e., Labyrinthine hydrops", "Definition", "Inner ear."],
    )
    parsed = parse_mcp_docx(path)
    assert parsed.display_name == "Ménière's Disease"
    chunks = build_chunks(parsed)
    assert chunks[0]["content"].startswith("Ménière's Disease — definition:")


def test_title_that_looks_like_marker_becomes_display_name(tmp_path):
    # The first non-empty paragraph is always the display name, even when it
    # would otherwise match a section marker.
    path = _write_docx(
        tmp_path,
        "MC900 Odd.docx",
        ["Signs and Symptoms Syndrome", "AKA i.e., SSS", "Definition", "Defined here."],
    )
    parsed = parse_mcp_docx(path)
    assert parsed.display_name == "Signs and Symptoms Syndrome"
    assert parsed.aliases == ["SSS"]
    assert parsed.sections == {"definition": ["Defined here."]}


def test_leading_blank_paragraphs_skipped_before_title(tmp_path):
    path = _write_docx(tmp_path, "MC011 Blanks.docx", ["", "", "Real Title", "Definition", "Text."])
    parsed = parse_mcp_docx(path)
    assert parsed.display_name == "Real Title"


def test_second_aka_line_before_markers_overwrites_first(tmp_path):
    # NOTE(potential-bug): a second preamble "AKA" line replaces (not extends)
    # the aliases from the first one.
    path = _write_docx(
        tmp_path,
        "MC012 TwoAka.docx",
        ["Title", "AKA i.e., First", "AKA i.e., Second", "Definition", "Text."],
    )
    parsed = parse_mcp_docx(path)
    assert parsed.aliases == ["Second"]


# --------------------------------------------------------------------------- #
# parse_mcp_docx: section routing
# --------------------------------------------------------------------------- #


def test_content_lands_in_right_sections(tmp_path):
    parsed = parse_mcp_docx(_write_corpus_docx(tmp_path))
    assert parsed.sections["definition"] == ["A sustained elevation of arterial blood pressure."]
    assert parsed.sections["prevalence"] == ["Common in adults worldwide."]
    assert parsed.sections["diagnosis"] == [
        "Repeated readings above the threshold on separate days."
    ]
    assert parsed.sections["suggestions"] == ["Discuss readings with a clinician."]


def test_structural_headers_hold_no_content(tmp_path):
    parsed = parse_mcp_docx(_write_corpus_docx(tmp_path))
    # No section key is ever a structural ("_"-prefixed) name.
    assert all(not key.startswith("_") for key in parsed.sections)
    # Text following "Introduction to MCP" is dropped entirely.
    all_lines = [ln for lines in parsed.sections.values() for ln in lines]
    assert "This intro text belongs to a structural header." not in all_lines


def test_unknown_paragraphs_before_first_marker_ignored(tmp_path):
    path = _write_docx(
        tmp_path,
        "MC020 Preamble.docx",
        ["Title", "AKA i.e., T", "Stray preamble note.", "Definition", "Actual text."],
    )
    parsed = parse_mcp_docx(path)
    all_lines = [ln for lines in parsed.sections.values() for ln in lines]
    assert "Stray preamble note." not in all_lines
    assert parsed.sections["definition"] == ["Actual text."]


def test_text_between_symptoms_marker_and_table_kept_in_order(tmp_path):
    parsed = parse_mcp_docx(_write_corpus_docx(tmp_path))
    assert parsed.sections["symptoms"] == [
        "Often silent in early stages.",
        "Symptom: Headache; Type: Common; Note: Morning pattern",
        "Symptom: Dizziness; Type: Occasional",
    ]


def test_risk_profiles_table_flattened(tmp_path):
    parsed = parse_mcp_docx(_write_corpus_docx(tmp_path))
    assert parsed.sections["risk_profiles"] == [
        "Profile: Sedentary adults; Reason: Low activity",
        "Profile: Smokers; Reason: Vascular strain",
    ]


def test_marker_matching_case_insensitive_startswith(tmp_path):
    path = _write_docx(
        tmp_path,
        "MC021 Caps.docx",
        ["Title", "DEFINITION and general overview", "Body text.", "PREVALENCE data", "Rates."],
    )
    parsed = parse_mcp_docx(path)
    assert parsed.sections["definition"] == ["Body text."]
    assert parsed.sections["prevalence"] == ["Rates."]


def test_marker_paragraph_itself_not_stored_as_content(tmp_path):
    parsed = parse_mcp_docx(_write_corpus_docx(tmp_path))
    all_lines = [ln for lines in parsed.sections.values() for ln in lines]
    assert "Definition" not in all_lines
    assert "Symptoms" not in all_lines


def test_empty_section_present_with_no_lines(tmp_path):
    path = _write_docx(
        tmp_path,
        "MC022 Empty.docx",
        ["Title", "Signs", "Suggestions", "See a clinician."],
    )
    parsed = parse_mcp_docx(path)
    assert parsed.sections["signs"] == []
    assert parsed.sections["suggestions"] == ["See a clinician."]


def test_table_before_any_marker_ignored(tmp_path):
    path = _write_docx(
        tmp_path,
        "MC023 EarlyTable.docx",
        ["Title", "AKA i.e., T", "Definition", "Text."],
        tables_after={1: [["H1", "H2"], ["a", "b"]]},  # table right after AKA line
    )
    parsed = parse_mcp_docx(path)
    all_lines = [ln for lines in parsed.sections.values() for ln in lines]
    assert all("H1" not in ln and "H2" not in ln for ln in all_lines)
    assert parsed.sections["definition"] == ["Text."]


def test_table_after_structural_header_ignored(tmp_path):
    path = _write_docx(
        tmp_path,
        "MC024 StructTable.docx",
        ["Title", "Definition", "Body.", "Research Draft of MCP"],
        tables_after={3: [["H", "H2"], ["CELLA", "CELLB"]]},
    )
    parsed = parse_mcp_docx(path)
    all_lines = [ln for lines in parsed.sections.values() for ln in lines]
    assert all("CELLA" not in ln and "CELLB" not in ln for ln in all_lines)
    assert parsed.sections["definition"] == ["Body."]


def test_aka_paragraph_after_first_marker_is_plain_content(tmp_path):
    path = _write_docx(
        tmp_path,
        "MC025 LateAka.docx",
        ["Title", "AKA i.e., Real Alias", "Definition", "AKA something in body text."],
    )
    parsed = parse_mcp_docx(path)
    assert parsed.aliases == ["Real Alias"]
    assert parsed.sections["definition"] == ["AKA something in body text."]


def test_content_paragraph_starting_with_marker_word_switches_section(tmp_path):
    # NOTE(potential-bug): marker matching is plain startswith on any
    # paragraph, so a body sentence beginning with a marker word ("Signs of
    # trouble...") is treated as a new section marker and re-routes content.
    path = _write_docx(
        tmp_path,
        "MC026 Hijack.docx",
        ["Title", "Definition", "Signs of trouble appear early.", "More definition text."],
    )
    parsed = parse_mcp_docx(path)
    assert parsed.sections["definition"] == []
    assert parsed.sections["signs"] == ["More definition text."]


def test_all_twenty_markers_route_or_drop(tmp_path):
    # Every marker either becomes a section (with its one content line) or, if
    # structural, drops its content.
    markers = list(SECTION_MARKERS)
    paragraphs = ["Full Coverage Condition", "AKA i.e., FCC"]
    for i, marker in enumerate(markers):
        paragraphs.append(marker.title())
        paragraphs.append(f"content-{i}")
    path = _write_docx(tmp_path, "MC030 Full.docx", paragraphs)
    parsed = parse_mcp_docx(path)
    for i, marker in enumerate(markers):
        section = SECTION_MARKERS[marker]
        if section.startswith("_"):
            assert section not in parsed.sections
        else:
            assert parsed.sections[section] == [f"content-{i}"]


# --------------------------------------------------------------------------- #
# _table_to_lines: flattening
# --------------------------------------------------------------------------- #


def test_table_with_zero_rows(tmp_path):
    assert _table_to_lines(_table([], cols=2)) == []


def test_table_header_only(tmp_path):
    assert _table_to_lines(_table([["Symptom", "Type", "Note"]])) == []


def test_table_basic_rows():
    t = _table(
        [
            ["Symptom", "Type", "Note"],
            ["Fever", "Acute", "Mild"],
            ["Cough", "Chronic", "Dry"],
        ]
    )
    assert _table_to_lines(t) == [
        "Symptom: Fever; Type: Acute; Note: Mild",
        "Symptom: Cough; Type: Chronic; Note: Dry",
    ]


def test_table_empty_cells_skipped():
    t = _table([["Symptom", "Type", "Note"], ["Fever", "", "Mild"]])
    assert _table_to_lines(t) == ["Symptom: Fever; Note: Mild"]


def test_table_whitespace_only_cell_skipped():
    t = _table([["Symptom", "Note"], ["Fever", "   "]])
    assert _table_to_lines(t) == ["Symptom: Fever"]


def test_table_all_empty_data_row_emits_no_line():
    t = _table([["Symptom", "Note"], ["", ""], ["Fever", "Mild"]])
    assert _table_to_lines(t) == ["Symptom: Fever; Note: Mild"]


def test_table_duplicate_header_cell_pairs_collapse():
    # Same (header, cell) pair repeated across columns collapses to one part —
    # the shape produced by merged cells under merged headers.
    t = _table([["Symptom", "Symptom"], ["Fever", "Fever"]])
    assert _table_to_lines(t) == ["Symptom: Fever"]


def test_table_merged_data_cells_collapse_when_headers_match():
    doc = Document()
    t = doc.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "Symptom"
    t.cell(0, 1).text = "Symptom"
    merged = t.cell(1, 0).merge(t.cell(1, 1))
    merged.text = "Fever"
    # A merged cell is reported once per spanned grid column.
    assert [c.text for c in t.rows[1].cells] == ["Fever", "Fever"]
    assert _table_to_lines(t) == ["Symptom: Fever"]


def test_table_merged_data_cells_under_different_headers_repeat():
    # NOTE(potential-bug): dedupe keys on (header, cell), so a data cell
    # merged across columns with *different* headers is emitted twice.
    doc = Document()
    t = doc.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "Type"
    t.cell(0, 1).text = "Note"
    merged = t.cell(1, 0).merge(t.cell(1, 1))
    merged.text = "Shared"
    assert _table_to_lines(t) == ["Type: Shared; Note: Shared"]


def test_table_multiline_cell_joined_with_slashes():
    doc = Document()
    t = doc.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "Symptom"
    t.cell(0, 1).text = "Note"
    t.cell(1, 0).text = "Itching"
    cell = t.cell(1, 1)
    cell.text = "worse at night"
    cell.add_paragraph("improves with cream")
    assert cell.text == "worse at night\nimproves with cream"
    assert _table_to_lines(t) == ["Symptom: Itching; Note: worse at night / improves with cream"]


def test_table_empty_header_uses_bare_cell_text():
    t = _table([["", "Note"], ["standalone", "n1"]])
    assert _table_to_lines(t) == ["standalone; Note: n1"]


def test_table_unicode_cells():
    t = _table([["Symptom", "Note"], ["Vertige — rotatoire", "Crises de Ménière"]])
    assert _table_to_lines(t) == ["Symptom: Vertige — rotatoire; Note: Crises de Ménière"]


def test_table_roundtrips_through_saved_docx(tmp_path):
    # Same flattening when the table goes through a real save/load cycle.
    path = _write_docx(
        tmp_path,
        "MC040 Tbl.docx",
        ["Title", "Signs"],
        tables_after={1: [["Sign", "Note"], ["Pallor", "Marked"]]},
    )
    parsed = parse_mcp_docx(path)
    assert parsed.sections["signs"] == ["Sign: Pallor; Note: Marked"]


# --------------------------------------------------------------------------- #
# build_chunks
# --------------------------------------------------------------------------- #


def _parsed(sections, code="MC001", display="Hypertension", source="MC001 Hypertension.docx"):
    return ParsedMcp(
        code=code, display_name=display, aliases=[], sections=sections, source_file=source
    )


def test_build_chunks_empty_sections():
    assert build_chunks(_parsed({})) == []


def test_build_chunks_empty_section_produces_no_chunk():
    assert build_chunks(_parsed({"signs": []})) == []


def test_build_chunks_whitespace_only_lines_produce_no_chunk():
    assert build_chunks(_parsed({"signs": ["", "   ", "\t"]})) == []


def test_build_chunks_single_section_shape():
    chunks = build_chunks(_parsed({"definition": ["Line one.", "Line two."]}))
    assert len(chunks) == 1
    chunk = chunks[0]
    assert chunk["condition_code"] == "MC001"
    assert chunk["chunk_type"] == "definition"
    assert chunk["content"] == "Hypertension — definition:\nLine one.\nLine two."
    assert chunk["metadata"] == {
        "source": "mcp_master_profile",
        "source_file": "MC001 Hypertension.docx",
        "display_name": "Hypertension",
        "section": "definition",
        "part": 1,
    }


def test_build_chunks_underscored_section_name_prettified_in_header():
    chunks = build_chunks(_parsed({"risk_profiles": ["Profile: X"]}))
    assert chunks[0]["content"].startswith("Hypertension — risk profiles:")
    # chunk_type and metadata keep the raw section name.
    assert chunks[0]["chunk_type"] == "risk_profiles"
    assert chunks[0]["metadata"]["section"] == "risk_profiles"


def test_build_chunks_every_chunk_starts_with_display_and_section(tmp_path):
    parsed = parse_mcp_docx(_write_corpus_docx(tmp_path))
    chunks = build_chunks(parsed)
    assert chunks  # corpus doc yields content
    for chunk in chunks:
        header, _, _ = chunk["content"].partition("\n")
        section_pretty = chunk["metadata"]["section"].replace("_", " ")
        assert header == f"Hypertension — {section_pretty}:"


def test_build_chunks_metadata_keys_exact(tmp_path):
    parsed = parse_mcp_docx(_write_corpus_docx(tmp_path))
    for chunk in build_chunks(parsed):
        assert set(chunk) == {"condition_code", "chunk_type", "content", "metadata"}
        assert set(chunk["metadata"]) == {
            "source",
            "source_file",
            "display_name",
            "section",
            "part",
        }
        assert chunk["metadata"]["source"] == "mcp_master_profile"


def test_build_chunks_splits_long_section_with_part_numbering():
    lines = [f"Row {i:03d}: " + "x" * 90 for i in range(60)]  # far beyond one chunk
    chunks = build_chunks(_parsed({"symptoms": lines}))
    assert len(chunks) > 1
    assert chunks[0]["chunk_type"] == "symptoms"
    for i, chunk in enumerate(chunks):
        if i > 0:
            assert chunk["chunk_type"] == f"symptoms_{i + 1}"
        assert chunk["metadata"]["part"] == i + 1
        assert chunk["metadata"]["section"] == "symptoms"
        assert len(chunk["content"]) <= MAX_CHUNK_CHARS
        assert chunk["content"].startswith("Hypertension — symptoms:\n")
    # All lines survive, in order, exactly once.
    rebuilt = []
    for chunk in chunks:
        rebuilt.extend(chunk["content"].split("\n")[1:])
    assert rebuilt == lines


def test_build_chunks_single_overlong_line_still_emitted():
    long_line = "y" * (MAX_CHUNK_CHARS + 700)
    chunks = build_chunks(_parsed({"definition": [long_line]}))
    assert len(chunks) == 1
    assert chunks[0]["chunk_type"] == "definition"
    assert chunks[0]["content"] == "Hypertension — definition:\n" + long_line
    assert len(chunks[0]["content"]) > MAX_CHUNK_CHARS


def test_build_chunks_overlong_line_flushes_previous_buffer():
    long_line = "z" * (MAX_CHUNK_CHARS + 1)
    chunks = build_chunks(_parsed({"definition": ["short line", long_line]}))
    assert [c["chunk_type"] for c in chunks] == ["definition", "definition_2"]
    assert chunks[0]["content"] == "Hypertension — definition:\nshort line"
    assert chunks[1]["content"] == "Hypertension — definition:\n" + long_line


def test_build_chunks_sections_emitted_in_insertion_order():
    chunks = build_chunks(_parsed({"definition": ["a"], "prevalence": ["b"], "signs": ["c"]}))
    assert [c["chunk_type"] for c in chunks] == ["definition", "prevalence", "signs"]


def test_split_to_chunks_boundary_exact_fit_not_split():
    header = "H" * 10
    line_a = "a" * 1000
    # size after line_a = 10 + 1001 = 1011; adding 788 + 1 = 1800 exactly: kept.
    line_b = "b" * 788
    chunks = _split_to_chunks(header, [line_a, line_b])
    assert len(chunks) == 1
    assert len(chunks[0]) == MAX_CHUNK_CHARS


def test_split_to_chunks_boundary_one_over_splits():
    header = "H" * 10
    line_a = "a" * 1000
    line_b = "b" * 789  # 1011 + 790 = 1801 > 1800: split
    chunks = _split_to_chunks(header, [line_a, line_b])
    assert len(chunks) == 2
    assert chunks[0] == header + "\n" + line_a
    assert chunks[1] == header + "\n" + line_b


def test_split_to_chunks_empty_lines_list():
    assert _split_to_chunks("Header:", []) == []


# --------------------------------------------------------------------------- #
# End-to-end: corpus doc -> chunks
# --------------------------------------------------------------------------- #


def test_end_to_end_corpus_chunks(tmp_path):
    parsed = parse_mcp_docx(_write_corpus_docx(tmp_path))
    chunks = build_chunks(parsed)
    types = [c["chunk_type"] for c in chunks]
    assert types == [
        "definition",
        "prevalence",
        "diagnosis",
        "risk_profiles",
        "symptoms",
        "suggestions",
    ]
    assert all(c["condition_code"] == "MC001" for c in chunks)
    symptoms = next(c for c in chunks if c["chunk_type"] == "symptoms")
    assert symptoms["content"] == (
        "Hypertension — symptoms:\n"
        "Often silent in early stages.\n"
        "Symptom: Headache; Type: Common; Note: Morning pattern\n"
        "Symptom: Dizziness; Type: Occasional"
    )
